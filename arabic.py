import fitz  # PyMuPDF
import docx
from PyQt5.QtWidgets import (
    QApplication, QLabel, QVBoxLayout, QWidget, QScrollArea, QPushButton,
    QFileDialog, QHBoxLayout, QTextEdit, QMessageBox
)
from PyQt5.QtGui import QPixmap
import sys
import os
import re
from bidi.algorithm import get_display
import arabic_reshaper
from docx.oxml.ns import qn
import pytesseract
from PIL import Image

# Configure pytesseract path for Windows users
pytesseract.pytesseract.tesseract_cmd = r'C:\\Program Files\\Tesseract-OCR\\tesseract.exe'

def sanitize_text(text):
    """Remove non-XML-compatible characters from text."""
    return re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

def reshape_arabic_text(text):
    """Reshape Arabic text for proper display and Word compatibility."""
    try:
        reshaped_text = arabic_reshaper.reshape(text)
        bidi_text = get_display(reshaped_text)
        return bidi_text
    except Exception as e:
        print(f"Error reshaping Arabic text: {e}")
        return text

def ocr_arabic_from_image(image):
    """Extract Arabic text from an image using OCR."""
    try:
        custom_config = '--oem 3 --psm 6 -l ara'
        text = pytesseract.image_to_string(image, config=custom_config)
        return reshape_arabic_text(sanitize_text(text))
    except Exception as e:
        print(f"Error in OCR: {e}")
        return ""

def pdf_to_word(pdf_path, output_word_path):
    """Convert a PDF file to a Word document with improved Arabic text and OCR support."""
    try:
        with fitz.open(pdf_path) as pdf_document:
            doc = docx.Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            for page_num in range(len(pdf_document)):
                page = pdf_document.load_page(page_num)
                text = page.get_text("text")
                if not text.strip():
                    pix = page.get_pixmap()
                    image_path = f"temp_ocr_page_{page_num + 1}.png"
                    pix.save(image_path)
                    text = ocr_arabic_from_image(Image.open(image_path))
                    os.remove(image_path)
                else:
                    text = reshape_arabic_text(sanitize_text(text))

                if text.strip():
                    for line in text.split('\n'):
                        paragraph = doc.add_paragraph()
                        run = paragraph.add_run(line)
                        run.font.name = 'Arial'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

            doc.save(output_word_path)
            print(f"PDF converted to Word and saved as '{output_word_path}'.")
    except Exception as e:
        print(f"Error in pdf_to_word: {e}")

def save_edited_text_to_pdf(original_pdf_path, edited_text):
    """Save edited text back to a new PDF file."""
    try:
        edited_pdf_path = original_pdf_path.replace(".pdf", "_edited.pdf")
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), edited_text, fontsize=12)
        doc.save(edited_pdf_path)
        doc.close()
        return edited_pdf_path
    except Exception as e:
        print(f"Error saving edited PDF: {e}")
        return None

class PDFEditorApp(QWidget):
    """Main application window for the PDF Editor and Viewer."""

    def __init__(self):
        super().__init__()
        self.setWindowTitle("PDF Editor & Viewer with Arabic OCR Support")
        self.resize(900, 700)
        self.pdf_path = ""
        self.init_ui()

    def init_ui(self):
        """Initialize the GUI layout and components."""
        layout = QVBoxLayout()
        button_layout = QHBoxLayout()

        self.select_button = QPushButton("Select PDF")
        self.select_button.clicked.connect(self.select_pdf)

        self.convert_button = QPushButton("Convert to Word (Arabic + OCR Support)")
        self.convert_button.clicked.connect(self.convert_pdf_to_word)
        self.convert_button.setEnabled(False)

        self.view_button = QPushButton("View PDF")
        self.view_button.clicked.connect(self.view_pdf)
        self.view_button.setEnabled(False)

        self.edit_button = QPushButton("Edit & Save PDF")
        self.edit_button.clicked.connect(self.edit_pdf)
        self.edit_button.setEnabled(False)

        self.text_editor = QTextEdit()
        self.text_editor.setPlaceholderText("Edit PDF text here...")
        self.text_editor.setFixedHeight(200)

        button_layout.addWidget(self.select_button)
        button_layout.addWidget(self.view_button)
        button_layout.addWidget(self.convert_button)
        button_layout.addWidget(self.edit_button)

        self.scroll_area = QScrollArea()
        self.scroll_widget = QWidget()
        self.scroll_layout = QVBoxLayout(self.scroll_widget)
        self.scroll_area.setWidgetResizable(True)
        self.scroll_area.setWidget(self.scroll_widget)

        layout.addLayout(button_layout)
        layout.addWidget(self.scroll_area)
        layout.addWidget(self.text_editor)
        self.setLayout(layout)

    def select_pdf(self):
        """Handle PDF file selection."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Select PDF File", "", "PDF Files (*.pdf)")
        if file_path:
            self.pdf_path = file_path
            self.view_button.setEnabled(True)
            self.convert_button.setEnabled(True)
            self.edit_button.setEnabled(True)
            self.load_pdf_text()

    def load_pdf_text(self):
        """Load PDF text into the editor with Arabic OCR support."""
        try:
            with fitz.open(self.pdf_path) as pdf_document:
                all_text = ""
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    text = page.get_text("text")
                    if not text.strip():
                        pix = page.get_pixmap()
                        image_path = f"temp_ocr_page_{page_num + 1}.png"
                        pix.save(image_path)
                        text = ocr_arabic_from_image(Image.open(image_path))
                        os.remove(image_path)
                    else:
                        text = reshape_arabic_text(sanitize_text(text))
                    all_text += text + "\n"
                self.text_editor.setPlainText(all_text)
        except Exception as e:
            print(f"Error loading PDF text: {e}")

    def view_pdf(self):
        """Render and display PDF pages as images."""
        try:
            for i in reversed(range(self.scroll_layout.count())):
                widget = self.scroll_layout.itemAt(i).widget()
                if widget:
                    widget.setParent(None)

            with fitz.open(self.pdf_path) as pdf_document:
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    pix = page.get_pixmap()
                    image_path = f"temp_page_{page_num + 1}.png"
                    pix.save(image_path)
                    label = QLabel()
                    label.setPixmap(QPixmap(image_path))
                    self.scroll_layout.addWidget(label)
                    os.remove(image_path)
        except Exception as e:
            print(f"Error in view_pdf: {e}")

    def convert_pdf_to_word(self):
        """Convert the loaded PDF to a Word document with Arabic and OCR support."""
        try:
            output_word_path, _ = QFileDialog.getSaveFileName(
                self, "Save Word File", "output.docx", "Word Files (*.docx)"
            )
            if output_word_path:
                pdf_to_word(self.pdf_path, output_word_path)
                QMessageBox.information(self, "Success", f"PDF converted to Word at:\n{output_word_path}")
        except Exception as e:
            print(f"Error in convert_pdf_to_word: {e}")

    def edit_pdf(self):
        """Save the edited text to a new PDF file."""
        edited_text = self.text_editor.toPlainText()
        edited_path = save_edited_text_to_pdf(self.pdf_path, edited_text)
        if edited_path:
            QMessageBox.information(self, "Success", f"Edited PDF saved at:\n{edited_path}")

if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = PDFEditorApp()
    window.show()
    sys.exit(app.exec_())